import os
import spacy
import docx
import shutil
import tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import streamlit as st

# Load NLP model
nlp = spacy.load("en_core_web_sm")

def extract_experience_and_education(resume_text):
    sections = {"experience": [], "education": []}
    lines = resume_text.splitlines()
    current_section = None

    for line in lines:
        line_lower = line.strip().lower()
        if "experience" in line_lower:
            current_section = "experience"
            continue
        elif "education" in line_lower:
            current_section = "education"
            continue
        elif line.strip() == "" or any(word in line_lower for word in ["summary", "skills", "project"]):
            current_section = None

        if current_section:
            sections[current_section].append(line.strip())

    return sections

def extract_name_role(text):
    lines = text.strip().split("\n")
    name = lines[0] if lines else "Your Name"
    role = lines[1] if len(lines) > 1 else "Your Role"
    return name.strip(), role.strip()

def extract_contact(text):
    email = phone = github = ""
    for line in text.splitlines():
        if "@" in line:
            email = line.strip()
        elif any(c.isdigit() for c in line) and len(line.strip()) > 8:
            phone = line.strip()
        elif "github.com" in line:
            github = line.strip()
    return email, phone, github

def write_heading(doc, text, size=14, bold=True):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para

def add_circular_image(doc, image_path):
    # Adds image and moves cursor to the next line
    if image_path:
        doc.add_picture(image_path, width=Inches(1.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def create_resume(resume_text, jd_text, image_path=None):
    doc = Document()

    name, role = extract_name_role(resume_text)
    email, phone, github = extract_contact(resume_text)
    experience_data = extract_experience_and_education(resume_text)

    # Extract JD skills
    jd_keywords = [token.text.lower() for token in nlp(jd_text) if token.is_alpha and not token.is_stop]
    jd_keywords = sorted(set(jd_keywords))

    # Header
    if image_path:
        add_circular_image(doc, image_path)
    doc.add_heading(name, 0)
    doc.add_paragraph(role)
    doc.add_paragraph(f"📧 {email} | 📞 {phone} | 🌐 {github}")
    doc.add_paragraph("_" * 70)

    # Summary
    write_heading(doc, "Professional Summary")
    doc.add_paragraph(f"Experienced professional with proven expertise in {', '.join(jd_keywords[:8])}. Passionate about delivering quality solutions.")

    # Skills
    write_heading(doc, "Skills")
    sorted_skills = sorted(jd_keywords)
    for skill in sorted_skills:
        doc.add_paragraph(f"\u2022 {skill.title()}", style='List Bullet')

    # Experience
    if experience_data['experience']:
        write_heading(doc, "Experience")
        for item in experience_data['experience']:
            doc.add_paragraph(item)

    # Education
    if experience_data['education']:
        write_heading(doc, "Education")
        for item in experience_data['education']:
            doc.add_paragraph(item)

    # Save to BytesIO
    temp_buffer = BytesIO()
    doc.save(temp_buffer)
    temp_buffer.seek(0)
    return temp_buffer

# ===== Streamlit UI =====
st.title("📄 Smart Resume Generator")
resume_input = st.text_area("Paste your Resume Text")
jd_input = st.text_area("Paste the Job Description")
image_input = st.file_uploader("Optional: Upload your Profile Image (JPG/PNG)", type=["jpg", "png"])

if st.button("Generate Optimized Resume"):
    if resume_input and jd_input:
        with st.spinner("Building resume..."):
            img_path = None
            if image_input:
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                tmp.write(image_input.read())
                tmp.close()
                img_path = tmp.name

            result = create_resume(resume_input, jd_input, img_path)

            st.download_button(
                label="📥 Download Resume",
                data=result,
                file_name="Optimized_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            if img_path:
                os.remove(img_path)
    else:
        st.warning("Please provide both resume text and job description.")